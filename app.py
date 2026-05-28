"""Resume Reader Application

This script reads and displays the content of a Word document (.docx) resume file.
"""

import sys
from pathlib import Path
from docx import Document


def read_resume(file_path):
    """
    Read a resume from a .docx file and display its content.
    
    Args:
        file_path (str): Path to the .docx resume file
    """
    file_path = Path(file_path)
    
    # Validate file exists
    if not file_path.exists():
        print(f"Error: File '{file_path}' not found.")
        sys.exit(1)
    
    # Validate file extension
    if file_path.suffix.lower() != '.docx':
        print(f"Error: File must be a .docx file. Got: {file_path.suffix}")
        sys.exit(1)
    
    try:
        # Load the document
        doc = Document(file_path)
        
        print(f"\n{'='*60}")
        print(f"Resume: {file_path.name}")
        print(f"{'='*60}\n")
        
        # Extract and display all paragraphs
        if not doc.paragraphs:
            print("No content found in the resume.")
            return
        
        for para in doc.paragraphs:
            if para.text.strip():  # Only print non-empty paragraphs
                print(para.text)
        
        # Display tables if any
        if doc.tables:
            print(f"\n{'-'*60}")
            print("Tables in document:")
            print(f"{'-'*60}\n")
            
            for table_idx, table in enumerate(doc.tables, 1):
                print(f"Table {table_idx}:")
                for row in table.rows:
                    print(" | ".join(cell.text.strip() for cell in row.cells))
                print()
        
        print(f"{'='*60}\n")
        
    except Exception as e:
        print(f"Error reading file: {e}")
        sys.exit(1)


def main():
    """Main entry point."""
    if len(sys.argv) != 2:
        print("Usage: python app.py <path_to_resume.docx>")
        print("\nExample: python app.py my_resume.docx")
        sys.exit(1)
    
    resume_path = sys.argv[1]
    read_resume(resume_path)


if __name__ == "__main__":
    main()
